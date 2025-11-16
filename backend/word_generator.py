from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os

class WordMeetingGenerator:
    def __init__(self):
        """初始化Word文档生成器"""
        self.document = None
    
    def _setup_document_styles(self):
        """设置文档样式"""
        # 获取样式集合
        styles = self.document.styles
        
        # 设置正文样式
        try:
            normal_style = styles['Normal']
            font = normal_style.font
            font.name = '宋体'
            font.size = Pt(12)
            
            # 设置段落格式
            paragraph_format = normal_style.paragraph_format
            paragraph_format.line_spacing = 1.15
            paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
    
    def create_meeting_document(self, meeting_info, output_path):
        """
        创建会议记录文档（表格形式）
        
        :param meeting_info: 会议信息字典
        :param output_path: 输出文件路径
        """
        # 每次创建文档时使用新的Document实例
        self.document = Document()
        self._setup_document_styles()
        # 设置文档标题
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run("会议记录")
        title_run.font.name = '微软雅黑'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 73, 125)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        self.document.add_paragraph()
        
        # 创建会议基本信息表格
        self._create_basic_info_table(meeting_info)
        
        # 创建会议内容记录表格
        self._create_content_table(meeting_info)
        
        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.document.save(output_path)
        
        return output_path
    
    def _create_basic_info_table(self, meeting_info):
        """
        创建会议基本信息表格
        
        :param meeting_info: 会议信息字典
        """
        # 创建3行2列的表格
        table = self.document.add_table(rows=3, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 设置表格宽度
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(8)  # 设置每个单元格宽度
        
        # 合并单元格
        # 会议主题行
        table.cell(0, 0).merge(table.cell(0, 1))
        
        # 设置标题行样式和内容
        theme_cell = table.cell(0, 0)
        theme_cell.text = "会议主题："
        
        # 添加主题内容
        theme_run = theme_cell.paragraphs[0].add_run(meeting_info.get('theme', '未指定'))
        theme_run.font.name = '宋体'
        theme_run.font.size = Pt(12)
        
        # 设置主持人行
        host_cell = table.cell(1, 0)
        host_cell.text = "主持人："
        
        # 添加主持人内容
        host_run = host_cell.paragraphs[0].add_run(meeting_info.get('host', '未指定'))
        host_run.font.name = '宋体'
        host_run.font.size = Pt(12)
        
        # 会议地点行 - 确保使用正确的提取信息
        location_cell = table.cell(2, 0)
        location_cell.text = "会议地点："
        
        # 添加会议地点内容
        location_value = meeting_info.get('location', '未指定')
        location_run = location_cell.paragraphs[0].add_run(location_value)
        location_run.font.name = '宋体'
        location_run.font.size = Pt(12)
        
        # 会议时长行 - 确保使用正确的提取信息
        duration_cell = table.cell(2, 1)
        duration_cell.text = "会议时长："
        
        # 添加会议时长内容
        duration_value = meeting_info.get('duration', '未指定')
        duration_run = duration_cell.paragraphs[0].add_run(duration_value)
        duration_run.font.name = '宋体'
        duration_run.font.size = Pt(12)
        
        # 添加空行
        self.document.add_paragraph()
        
        # 添加参会人员表格
        attendees_table = self.document.add_table(rows=1, cols=2)
        attendees_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 设置参会人员表格宽度
        attendees_table.cell(0, 0).width = Cm(4)
        attendees_table.cell(0, 1).width = Cm(12)
        
        # 设置参会人员标题
        attendees_title_cell = attendees_table.cell(0, 0)
        attendees_title_cell.text = "参会人员："
        
        # 添加参会人员内容 - 确保使用正确的提取信息
        attendees_value = meeting_info.get('attendees', '未指定')
        attendees_content_cell = attendees_table.cell(0, 1)
        attendees_content_cell.text = attendees_value
        
        # 设置字体
        for cell in attendees_table.columns[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
        
        for cell in attendees_table.columns[1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
        
        # 添加空行
        self.document.add_paragraph()
    
    def _create_content_table(self, meeting_info):
        """
        创建会议内容记录表格
        
        :param meeting_info: 会议信息字典
        """
        topics = meeting_info.get('topics', [])
        
        # 计算需要的行数：表头 + 每个议题 + 会前准备事项
        num_rows = 2 + len(topics)  # 1行表头 + 至少1行内容
        
        # 创建内容表格（2列：左侧为标签，右侧为内容）
        table = self.document.add_table(rows=num_rows, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 设置表格宽度
        table.cell(0, 0).merge(table.cell(0, 1))  # 合并表头
        
        # 设置表头样式和内容
        header_cell = table.cell(0, 0)
        header_paragraph = header_cell.paragraphs[0]
        header_run = header_paragraph.add_run("会议内容记录")
        header_run.font.name = '微软雅黑'
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(59, 89, 152)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置议题内容
        for i, topic in enumerate(topics, 1):
            # 当前议题的行索引
            topic_row_index = i
            
            # 议题内容单元格
            topic_cell = table.cell(topic_row_index, 0)
            topic_cell.merge(table.cell(topic_row_index, 1))  # 合并整行
            
            # 创建议题内容段落
            topic_para = topic_cell.paragraphs[0]
            topic_run = topic_para.add_run(f"议题{i}：{topic.get('topic', '')}")
            topic_run.font.name = '宋体'
            topic_run.font.size = Pt(12)
            topic_run.font.bold = True
            
            # 负责人信息（在同一单元格内添加新段落）- 移除英文标签，确保使用正确的提取信息
            leader_value = topic.get('leader', '未指定')
            leader_para = topic_cell.add_paragraph()
            leader_para.paragraph_format.left_indent = Cm(1.0)
            leader_run = leader_para.add_run(f"负责人：{leader_value}")
            leader_run.font.name = '宋体'
            leader_run.font.size = Pt(12)
            
            # 会前准备信息
            preparation_value = topic.get('preparation', '无')
            prep_para = topic_cell.add_paragraph()
            prep_para.paragraph_format.left_indent = Cm(1.0)
            prep_run = prep_para.add_run(f"会前准备：{preparation_value}")
            prep_run.font.name = '宋体'
            prep_run.font.size = Pt(12)
        
        # 添加会前准备事项（如果有）
        if meeting_info.get('preparation_items', '') and num_rows > 2:
            last_row_index = num_rows - 1
            
            # 合并最后一行
            prep_cell = table.cell(last_row_index, 0)
            prep_cell.merge(table.cell(last_row_index, 1))
            
            # 设置会前准备事项内容
            prep_para = prep_cell.paragraphs[0]
            prep_run = prep_para.add_run(f"会前准备事项：{meeting_info.get('preparation_items')}")
            prep_run.font.name = '宋体'
            prep_run.font.size = Pt(12)
    
    def _add_section(self, title, content):
        """添加一个文档章节（备用方法，当前不使用）"""
        # 章节标题
        heading_para = self.document.add_paragraph()
        heading_run = heading_para.add_run(title)
        heading_run.font.name = '微软雅黑'
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(59, 89, 152)
        heading_para.paragraph_format.space_before = Pt(12)
        heading_para.paragraph_format.space_after = Pt(6)
        
        # 章节内容
        content_para = self.document.add_paragraph()
        content_run = content_para.add_run(content)
        content_run.font.name = '宋体'
        content_run.font.size = Pt(12)
        content_para.paragraph_format.left_indent = Inches(0.25)